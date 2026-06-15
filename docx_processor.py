"""Чтение шаблона `.docx`, замена плейсхолдера и сборка итогового документа."""

# Импортируется deepcopy, чтобы копировать XML-форматирование run без связи с исходным элементом.
from copy import deepcopy

# Импортируется Path, чтобы явно работать с путями к шаблону и итоговому файлу.
from pathlib import Path

# Импортируются Callable и Iterable для аннотаций обратного вызова прогресса и обхода параграфов.
from typing import Callable, Iterable

# Импортируется Document, чтобы читать и создавать `.docx` через python-docx.
from docx import Document

# Импортируется Pt, чтобы применять выбранный пользователем размер шрифта к вставленному номеру.
from docx.shared import Pt

# Импортируется Run, чтобы создавать новый run внутри существующего параграфа в нужной позиции.
from docx.text.run import Run

# Импортируется OxmlElement, чтобы вставлять новый run на уровне XML без переноса текста в конец абзаца.
from docx.oxml import OxmlElement

# Импортируется qn, чтобы корректно задавать XML-атрибуты Word для разных групп шрифтов.
from docx.oxml.ns import qn

# Импортируется Composer, чтобы объединять несколько `.docx` в один итоговый документ.
from docxcompose.composer import Composer


# Строка плейсхолдера хранится в константе, чтобы все проверки и замены использовали один формат.
PLACEHOLDER = "[№]"

# Тип обратного вызова прогресса фиксирует формат: выполненные шаги и общее число шагов.
ProgressCallback = Callable[[int, int], None]


class DocxProcessingError(RuntimeError):
    """Ошибка обработки `.docx`, пригодная для показа пользователю.

    Класс принимает стандартный текст ошибки.
    Ничего не возвращает.
    Побочный эффект: останавливает создание итогового документа при проблеме с шаблоном или записью.
    """


def document_contains_placeholder(template_path: Path) -> bool:
    """Проверяет, есть ли в `.docx` хотя бы один плейсхолдер `[№]`.

    template_path: путь к файлу шаблона.
    Возвращает True, если плейсхолдер найден в тексте, таблицах или колонтитулах.
    Побочный эффект: читает `.docx` с диска и может выбросить DocxProcessingError при ошибке чтения.
    """
    try:
        # Документ открывается только для чтения и проверки содержимого.
        document = Document(str(template_path))
    except Exception as exc:
        raise DocxProcessingError(f"Не удалось открыть DOCX-шаблон: {exc}") from exc

    return _document_has_placeholder(document)


def create_numbered_document(
    template_path: Path,
    output_path: Path,
    numbers: list[str],
    font_name: str,
    font_size: int,
    progress_callback: ProgressCallback | None = None,
) -> None:
    """Создает итоговый `.docx` из нескольких копий шаблона.

    template_path: путь к исходному `.docx`-шаблону.
    output_path: путь, куда нужно сохранить итоговый `.docx`.
    numbers: список итоговых номеров, по одному номеру на копию шаблона.
    font_name: выбранное пользователем имя шрифта для вставленного номера.
    font_size: выбранный пользователем размер шрифта в пунктах.
    progress_callback: необязательная функция, которую вызывают после подготовки копии и сохранения.
    Ничего не возвращает.
    Побочный эффект: читает шаблон, создает копии документа и записывает итоговый файл на диск.
    """
    if not numbers:
        raise DocxProcessingError("Список номеров пуст, итоговый документ создать нельзя.")

    # Количество шагов включает каждую копию и отдельный шаг сохранения итогового файла.
    total_steps = len(numbers) + 1

    try:
        # Первая копия становится базовым документом, к которому Composer добавит остальные копии.
        master_document = _build_single_copy(template_path, numbers[0], font_name, font_size)
        _notify_progress(progress_callback, 1, total_steps)
    except Exception as exc:
        raise DocxProcessingError(f"Не удалось подготовить первую копию шаблона: {exc}") from exc

    # Composer сохраняет структуру документа лучше, чем ручное копирование параграфов и таблиц.
    composer = Composer(master_document)

    for completed_copies, number in enumerate(numbers[1:], start=2):
        try:
            # Каждая копия создается из исходного шаблона, чтобы номера и замены не влияли друг на друга.
            next_document = _build_single_copy(template_path, number, font_name, font_size)

            # Разрыв добавляется в начало следующей копии, чтобы Word не создавал отдельную пустую страницу
            # из-за самостоятельного абзаца-разрыва после предыдущей копии.
            _prepend_page_break_to_first_paragraph(next_document)

            composer.append(next_document)
            _notify_progress(progress_callback, completed_copies, total_steps)
        except Exception as exc:
            raise DocxProcessingError(f"Не удалось добавить копию с номером {number}: {exc}") from exc

    try:
        # Сохранение выполняется в самом конце, чтобы при ошибке не оставлять частично собранный файл.
        composer.save(str(output_path))
        _notify_progress(progress_callback, total_steps, total_steps)
    except Exception as exc:
        raise DocxProcessingError(f"Не удалось сохранить итоговый DOCX-файл: {exc}") from exc


def _notify_progress(progress_callback: ProgressCallback | None, completed_steps: int, total_steps: int) -> None:
    """Передает наружу состояние прогресса, если вызывающий код его запросил.

    progress_callback: функция обратного вызова или None.
    completed_steps: количество завершенных шагов.
    total_steps: общее количество шагов обработки DOCX.
    Ничего не возвращает.
    Побочный эффект: вызывает progress_callback, если он передан.
    """
    if progress_callback is not None:
        progress_callback(completed_steps, total_steps)


def _build_single_copy(template_path: Path, number: str, font_name: str, font_size: int):
    """Создает одну копию шаблона с замененным номером.

    template_path: путь к исходному шаблону.
    number: номер, который нужно вставить вместо всех `[№]` в этой копии.
    font_name: имя шрифта для вставленного номера.
    font_size: размер шрифта для вставленного номера.
    Возвращает объект Document с выполненными заменами.
    Побочный эффект: читает шаблон с диска.
    """
    # Каждый номер получает собственный экземпляр документа, чтобы повторные замены не смешивались.
    document = Document(str(template_path))

    # Счетчик замен нужен, чтобы отличить корректный шаблон от файла без плейсхолдера.
    replacement_count = replace_placeholder_in_document(document, number, font_name, font_size)

    if replacement_count == 0:
        raise DocxProcessingError(f"В шаблоне не найден плейсхолдер {PLACEHOLDER}.")

    return document


def _prepend_page_break_to_first_paragraph(document) -> None:
    """Добавляет разрыв страницы перед первым содержимым документа.

    document: объект python-docx Document для очередной копии шаблона.
    Ничего не возвращает.
    Побочный эффект: вставляет XML-run с page break в начало первого параграфа документа.
    """
    if not document.paragraphs:
        raise DocxProcessingError("В копии шаблона нет параграфа для вставки разрыва страницы.")

    # Первый параграф выбран потому, что текст после разрыва остается в этом же параграфе и не образует
    # самостоятельную пустую страницу между копиями.
    first_paragraph = document.paragraphs[0]

    # XML-run создается вручную, чтобы вставить разрыв до первого текста, а не отдельным абзацем.
    break_run = OxmlElement("w:r")
    break_element = OxmlElement("w:br")
    break_element.set(qn("w:type"), "page")
    break_run.append(break_element)

    # Если у параграфа есть свойства, новый run должен стоять после них, иначе XML Word будет некорректным.
    paragraph_properties = first_paragraph._p.pPr
    insert_index = 1 if paragraph_properties is not None else 0
    first_paragraph._p.insert(insert_index, break_run)


def replace_placeholder_in_document(document, number: str, font_name: str, font_size: int) -> int:
    """Заменяет все `[№]` в документе на один номер.

    document: объект python-docx Document.
    number: номер для вставки в текущую копию шаблона.
    font_name: имя шрифта для вставленного номера.
    font_size: размер шрифта для вставленного номера.
    Возвращает количество выполненных замен.
    Побочный эффект: изменяет переданный Document в памяти.
    """
    # Общий счетчик нужен для проверки, что хотя бы один плейсхолдер действительно был заменен.
    total_replacements = 0

    for paragraph in _iter_document_paragraphs(document):
        total_replacements += _replace_placeholder_in_paragraph(paragraph, number, font_name, font_size)

    return total_replacements


def _document_has_placeholder(document) -> bool:
    """Проверяет наличие плейсхолдера в уже открытом документе.

    document: объект python-docx Document.
    Возвращает True, если `[№]` найден.
    Побочных эффектов нет.
    """
    for paragraph in _iter_document_paragraphs(document):
        if PLACEHOLDER in _get_paragraph_text(paragraph):
            return True

    return False


def _iter_document_paragraphs(document) -> Iterable:
    """Обходит параграфы тела документа, таблиц, headers и footers.

    document: объект python-docx Document.
    Возвращает итератор параграфов.
    Побочных эффектов нет.
    """
    yield from _iter_container_paragraphs(document)

    for section in document.sections:
        # Колонтитулы обрабатываются отдельно, потому что они не входят в document.paragraphs.
        yield from _iter_container_paragraphs(section.header)
        yield from _iter_container_paragraphs(section.footer)


def _iter_container_paragraphs(container) -> Iterable:
    """Обходит параграфы контейнера python-docx и вложенных таблиц.

    container: Document, Header, Footer, Cell или другой объект с paragraphs/tables.
    Возвращает итератор найденных параграфов.
    Побочных эффектов нет.
    """
    for paragraph in container.paragraphs:
        yield paragraph

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def _replace_placeholder_in_paragraph(paragraph, number: str, font_name: str, font_size: int) -> int:
    """Заменяет все плейсхолдеры внутри одного параграфа.

    paragraph: параграф python-docx.
    number: номер для вставки.
    font_name: имя шрифта для вставленного номера.
    font_size: размер шрифта для вставленного номера.
    Возвращает количество замен в параграфе.
    Побочный эффект: меняет runs параграфа в памяти.
    """
    # Количество замен защищает цикл от лишней работы и возвращается вызывающему коду.
    replacement_count = 0

    while PLACEHOLDER in _get_paragraph_text(paragraph):
        # Карта символов показывает, в каком run и на какой позиции находится каждый символ параграфа.
        paragraph_text, character_map = _build_character_map(paragraph)

        # Позиция первого плейсхолдера заменяется, затем карта строится заново для следующих совпадений.
        start_position = paragraph_text.find(PLACEHOLDER)

        if start_position < 0:
            break

        _replace_single_placeholder(paragraph, character_map, start_position, number, font_name, font_size)
        replacement_count += 1

    return replacement_count


def _build_character_map(paragraph) -> tuple[str, list[tuple[int, int]]]:
    """Строит общий текст параграфа и карту символов к runs.

    paragraph: параграф python-docx.
    Возвращает пару `(текст, карта)`, где карта хранит `(индекс_run, индекс_символа_в_run)`.
    Побочных эффектов нет.
    """
    # Части текста объединяются в один список, чтобы корректно находить `[№]`, даже если Word разбил его.
    text_parts: list[str] = []

    # Карта связывает каждый символ общего текста с конкретным run.
    character_map: list[tuple[int, int]] = []

    for run_index, run in enumerate(paragraph.runs):
        for char_index, character in enumerate(run.text):
            text_parts.append(character)
            character_map.append((run_index, char_index))

    return "".join(text_parts), character_map


def _replace_single_placeholder(
    paragraph,
    character_map: list[tuple[int, int]],
    start_position: int,
    number: str,
    font_name: str,
    font_size: int,
) -> None:
    """Заменяет один найденный плейсхолдер внутри параграфа.

    paragraph: параграф python-docx.
    character_map: карта символов, построенная для текущего состояния параграфа.
    start_position: позиция начала `[№]` в общем тексте параграфа.
    number: номер для вставки.
    font_name: имя шрифта для вставленного номера.
    font_size: размер шрифта для вставленного номера.
    Ничего не возвращает.
    Побочный эффект: изменяет runs параграфа.
    """
    # Конечная позиция нужна, чтобы понять, в каком run заканчивается плейсхолдер.
    end_position = start_position + len(PLACEHOLDER) - 1

    # Индексы начала и конца связывают плейсхолдер с реальными runs Word-документа.
    start_run_index, start_char_index = character_map[start_position]
    end_run_index, end_char_index = character_map[end_position]

    # Список runs читается перед изменениями, потому что вставка новых runs меняет структуру параграфа.
    runs = paragraph.runs

    # Первый и последний run плейсхолдера могут совпадать или быть разными при разбиении Word.
    start_run = runs[start_run_index]
    end_run = runs[end_run_index]

    # Текст до плейсхолдера должен остаться в исходном run на прежнем месте.
    text_before = start_run.text[:start_char_index]

    # Текст после плейсхолдера должен остаться сразу после вставленного номера.
    text_after = end_run.text[end_char_index + 1 :]

    if start_run_index == end_run_index:
        # При плейсхолдере внутри одного run исходный run оставляет только текст до `[№]`.
        original_run = start_run
        original_run.text = text_before

        # Новый run с номером вставляется сразу после исходного run, а не в конец параграфа.
        number_run = _insert_text_run_after(original_run, number, original_run)
        _apply_number_font(number_run, font_name, font_size)

        if text_after:
            # Текст после плейсхолдера возвращается отдельным run с исходным форматированием.
            _insert_text_run_after(number_run, text_after, original_run)

        return

    # Для плейсхолдера, разбитого на несколько runs, начало и конец сохраняют свои окружающие части.
    start_run.text = text_before
    end_run.text = text_after

    for run_index in range(start_run_index + 1, end_run_index):
        # Средние runs содержали только часть плейсхолдера, поэтому их видимый текст очищается.
        runs[run_index].text = ""

    # Номер помещается между текстом до плейсхолдера и оставшейся частью последнего run.
    number_run = _insert_text_run_after(start_run, number, start_run)
    _apply_number_font(number_run, font_name, font_size)


def _insert_text_run_after(reference_run: Run, text: str, source_format_run: Run) -> Run:
    """Вставляет новый run после указанного run.

    reference_run: run, после которого нужно вставить новый элемент.
    text: текст нового run.
    source_format_run: run, из которого копируется исходное форматирование.
    Возвращает созданный Run.
    Побочный эффект: меняет XML-структуру параграфа.
    """
    # Новый XML-run создается вручную, чтобы сохранить точную позицию вставки в параграфе.
    new_run_element = OxmlElement("w:r")

    if source_format_run._r.rPr is not None:
        # Форматирование копируется глубоко, чтобы новый run не ссылался на XML исходного run.
        new_run_element.append(deepcopy(source_format_run._r.rPr))

    reference_run._r.addnext(new_run_element)

    # Объект Run нужен, чтобы дальше пользоваться обычным API python-docx.
    new_run = Run(new_run_element, reference_run._parent)
    new_run.text = text

    return new_run


def _apply_number_font(run: Run, font_name: str, font_size: int) -> None:
    """Применяет выбранные шрифт и размер только к вставленному номеру.

    run: run со вставленным номером.
    font_name: имя системного шрифта.
    font_size: размер шрифта в пунктах.
    Ничего не возвращает.
    Побочный эффект: меняет форматирование run.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # XML-настройка rFonts нужна, чтобы Word применял выбранный шрифт и к кириллическим символам.
    run_properties = run._r.get_or_add_rPr()

    # Элемент w:rFonts может уже существовать после копирования форматирования из исходного run.
    run_fonts = run_properties.find(qn("w:rFonts"))

    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)

    for attribute_name in ("ascii", "hAnsi", "cs", "eastAsia"):
        # Один и тот же шрифт задается для латиницы, кириллицы и восточноазиатского диапазона Word.
        run_fonts.set(qn(f"w:{attribute_name}"), font_name)


def _get_paragraph_text(paragraph) -> str:
    """Возвращает видимый текст параграфа по runs.

    paragraph: параграф python-docx.
    Возвращает объединенный текст всех runs.
    Побочных эффектов нет.
    """
    # paragraph.text также доступен, но явное объединение runs совпадает с картой символов замены.
    return "".join(run.text for run in paragraph.runs)
