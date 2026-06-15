"""Интеграционные тесты обработки DOCX-шаблонов."""

# Импортируется ZipFile, чтобы проверить XML итогового DOCX без запуска Word.
from zipfile import ZipFile

# Импортируется ElementTree, чтобы разобрать document.xml и найти page break.
from xml.etree import ElementTree as ET

# Импортируется Document, чтобы создавать временные DOCX-шаблоны и читать результат.
from docx import Document

# Импортируются функции обработки DOCX, которые должны сохранить позицию плейсхолдера.
from docx_processor import (
    PLACEHOLDER,
    create_numbered_document,
    document_contains_placeholder,
    replace_placeholder_in_document,
)


def test_replace_placeholder_in_paragraph_table_and_multiple_copies(tmp_path) -> None:
    """Проверяет замену в параграфе, таблице и нескольких копиях."""
    # Временный шаблон содержит плейсхолдер в обычном абзаце и ячейке таблицы.
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"

    document = Document()
    document.add_paragraph("Договор [№] от 01.01.2026")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Табличный номер [№]"

    document.save(str(template_path))

    assert document_contains_placeholder(template_path)

    create_numbered_document(
        template_path,
        output_path,
        ["№ABC-111", "№ABC-222"],
        "Arial",
        12,
    )

    result_document = Document(str(output_path))

    # Общий текст результата используется для проверки двух копий и отсутствия плейсхолдера.
    result_text = "\n".join(paragraph.text for paragraph in result_document.paragraphs)

    for result_table in result_document.tables:
        for row in result_table.rows:
            for cell in row.cells:
                result_text += "\n" + "\n".join(paragraph.text for paragraph in cell.paragraphs)

    assert PLACEHOLDER not in result_text
    assert "Договор №ABC-111 от 01.01.2026" in result_text
    assert "Договор №ABC-222 от 01.01.2026" in result_text
    assert "Табличный номер №ABC-111" in result_text
    assert "Табличный номер №ABC-222" in result_text


def test_replace_placeholder_split_across_runs_keeps_position() -> None:
    """Проверяет замену плейсхолдера, разбитого Word на несколько runs."""
    # Плейсхолдер намеренно разбит на три run, чтобы проверить сложный сценарий Word.
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("До [")
    paragraph.add_run("№")
    paragraph.add_run("] после")

    replacement_count = replace_placeholder_in_document(document, "№ABC-123", "Arial", 12)

    assert replacement_count == 1
    assert paragraph.text == "До №ABC-123 после"


def test_replacement_applies_font_only_to_inserted_number() -> None:
    """Проверяет применение выбранного шрифта к run со вставленным номером."""
    document = Document()
    paragraph = document.add_paragraph("A[№]B")

    replace_placeholder_in_document(document, "№ABC-123", "Arial", 14)

    number_runs = [run for run in paragraph.runs if run.text == "№ABC-123"]

    assert len(number_runs) == 1
    assert number_runs[0].font.name == "Arial"
    assert number_runs[0].font.size.pt == 14
    assert paragraph.text == "A№ABC-123B"


def test_page_break_is_not_created_as_empty_standalone_paragraph(tmp_path) -> None:
    """Проверяет, что разделитель копий не создает пустой абзац-страницу."""
    # Шаблон с одним абзацем помогает точно проверить место вставки page break.
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"

    document = Document()
    document.add_paragraph("Документ [№]")
    document.save(str(template_path))

    create_numbered_document(
        template_path,
        output_path,
        ["№ABC-001", "№ABC-002"],
        "Arial",
        12,
    )

    # Итоговый XML проверяется напрямую, потому что python-docx не показывает позицию page break в тексте.
    with ZipFile(output_path) as archive:
        xml = archive.read("word/document.xml")

    root = ET.fromstring(xml)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = root.findall(".//w:body/w:p", namespace)

    # В документе должен быть один разрыв между двумя копиями.
    paragraphs_with_page_break = [
        paragraph
        for paragraph in paragraphs
        if paragraph.findall('.//w:br[@w:type="page"]', namespace)
    ]

    assert len(paragraphs_with_page_break) == 1

    # В абзаце с разрывом должен быть текст второй копии, а не пустой самостоятельный абзац.
    break_paragraph_text = "".join(
        text_node.text or "" for text_node in paragraphs_with_page_break[0].findall(".//w:t", namespace)
    )

    assert "Документ №ABC-002" in break_paragraph_text


def test_create_numbered_document_reports_progress_for_copies_and_save(tmp_path) -> None:
    """Проверяет события прогресса для копий и финального сохранения."""
    # Простой шаблон позволяет проверить саму последовательность прогресса без лишней структуры DOCX.
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"

    document = Document()
    document.add_paragraph("Документ [№]")
    document.save(str(template_path))

    # Список событий имитирует данные, которые GUI получает для progressbar.
    progress_events: list[tuple[int, int]] = []

    create_numbered_document(
        template_path,
        output_path,
        ["№ABC-001", "№ABC-002"],
        "Arial",
        12,
        progress_callback=lambda completed, total: progress_events.append((completed, total)),
    )

    assert progress_events == [(1, 3), (2, 3), (3, 3)]
